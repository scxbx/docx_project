from docx import Document

# 读取文档
doc = Document('C:/Users/sc/Desktop/test_docx/sample.docx')  # filename为word文档

# 获取文档中的表格
doc.tables  # 获取文档的表格个数 len(doc.tables)

# 读取第1个表格
tb1 = doc.tables[0]

# 获取第一个表格的行
tb1.rows  # 获取表格的行数len(tb1.rows)

# 读取表格的第一行的单元格
row_cells = tb1.rows[0].cells

# 读取第一行所有单元格的内容
# for cell in row_cells:
#     print(cell.text)
print(tb1.cell(2, 1).text)
print(tb1.cell(20, 0).text)

headcount = 0


# to judge whether a string contains all keys in a list
def checkAllKeysInAString(my_list, my_str):
    for key in my_list:
        if key not in my_str:
            return False
    return True


for table in doc.tables:
    print(table.cell(0, 0).text)

print(len(doc.tables))

print(table.cell(26, 1).text)
